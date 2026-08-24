# doc_processor.py
"""
Multi-Format Sales Document Ingestion & Intelligent Strategy Chunker.
Supports PDF, DOCX, TXT, Markdown, and CSV files.
Splits custom sales playbooks into structured Battlecard Chunks with rich semantic metadata.
"""

import os
import io
import re
import csv
import logging
from typing import List, Dict, Any, Optional

logger = logging.getLogger("DocProcessor")

def sanitize_unicode(text: str) -> str:
    """Removes lone unicode surrogates and unencodable characters that break JSON serialization."""
    if not isinstance(text, str):
        return ""
    # Strip lone surrogates
    cleaned = "".join(ch for ch in text if not (0xD800 <= ord(ch) <= 0xDFFF))
    return cleaned.encode("utf-8", "ignore").decode("utf-8", "ignore")

class DocumentProcessor:
    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> str:
        """Extracts plain text from various file formats."""
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                pages = [p.extract_text() or "" for p in reader.pages]
                raw_text = "\n".join(pages).strip()
                return sanitize_unicode(raw_text)
            except Exception as e:
                logger.error(f"Error parsing PDF {filename}: {e}")
                raise ValueError(f"Failed to read PDF document: {e}")

        elif ext in (".docx", ".doc"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            full_text.append(" | ".join(row_text))
                return sanitize_unicode("\n".join(full_text).strip())
            except Exception as e:
                # Fallback to XML extraction from zip
                try:
                    import zipfile
                    import xml.etree.ElementTree as ET
                    with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                        xml_content = zf.read('word/document.xml')
                        tree = ET.fromstring(xml_content)
                        texts = [node.text for node in tree.iter() if node.text]
                        return sanitize_unicode(" ".join(texts).strip())
                except Exception as inner_e:
                    logger.error(f"Error parsing DOCX {filename}: {e} / {inner_e}")
                    raise ValueError(f"Failed to read Word document: {e}")

        elif ext in (".txt", ".md", ".json", ".log"):
            try:
                return sanitize_unicode(file_bytes.decode("utf-8", errors="ignore").strip())
            except Exception as e:
                logger.error(f"Error reading text file {filename}: {e}")
                raise ValueError(f"Failed to read text file: {e}")

        elif ext == ".csv":
            try:
                content = sanitize_unicode(file_bytes.decode("utf-8", errors="ignore"))
                reader = csv.reader(io.StringIO(content))
                rows = []
                for r in reader:
                    if any(r):
                        rows.append(" | ".join([c.strip() for c in r if c.strip()]))
                return "\n".join(rows).strip()
            except Exception as e:
                logger.error(f"Error reading CSV {filename}: {e}")
                raise ValueError(f"Failed to read CSV: {e}")

        else:
            # Attempt generic UTF-8 fallback
            try:
                return sanitize_unicode(file_bytes.decode("utf-8", errors="ignore").strip())
            except Exception:
                raise ValueError(f"Unsupported file format '{ext}'. Please upload PDF, DOCX, TXT, MD, or CSV.")

    @classmethod
    def chunk_strategies(cls, text: str, source_filename: str) -> List[Dict[str, Any]]:
        """
        Intelligently parses extracted text into high-resolution, granular strategy battlecards.
        Supports:
        1. Structured Q&A / Objection blocks (Q1., Question:, Objection:, Scenario:, Client says:)
        2. Numbered items & Rules (1., 2., 1), 2), Rule 1:, Strategy 1:, Tip 1:, Step 1:)
        3. Markdown Headers (# Header, ## Section)
        4. Bullet points & list items (•, -, *, ▪, ►)
        5. Semantic Sliding Paragraph / Sentence Windows (80-140 words per battlecard)
        """
        if not text or len(text.strip()) < 10:
            return []

        clean_text = sanitize_unicode(text.replace("\ufffd", "-").replace("\u00ad", ""))
        raw_lines = [l.strip() for l in clean_text.splitlines()]
        clean_text = "\n".join([l for l in raw_lines if l])

        cards = []

        # 1. Check for Structured Q&A / Objection Blocks
        split_pattern = r'(?=(?:^|\n)\s*(?:Q\d+\.?|Q\s*:|Question\b|Objection\b|Scenario\b|Client\s+(?:says|asks|demands|requires)\b))'
        raw_blocks = [b.strip() for b in re.split(split_pattern, clean_text, flags=re.IGNORECASE) if b.strip()]

        if len(raw_blocks) >= 2:
            card_id = 1
            for block in raw_blocks:
                if not re.match(r'^(?:Q\d+\.?|Q\s*:|Question\b|Objection\b|Scenario\b|Client\s+(?:says|asks|demands|requires)\b)', block, re.IGNORECASE):
                    continue

                sub_m = re.search(r'^(?:Q\d+\.?|Q\s*:|Question(?:\s*\d+)?:?|Objection(?:\s*\d+)?:?|Scenario(?:\s*\d+)?:?|Client\s+(?:says|asks|demands|requires):?)\s*(.+?)\n+(?:Context\s*/\s*Rationale|Context|Background|Why)\s*\n+(.+?)\n+(?:Exact\s*Strategy\s*/\s*Pitch|Pitch|Strategy|Response|Answer|Solution)\s*\n+(.+)', block, re.DOTALL | re.IGNORECASE)
                if sub_m:
                    q_text = sub_m.group(1).strip()
                    context_text = sub_m.group(2).strip()
                    pitch_text = sub_m.group(3).strip()
                else:
                    sub_m2 = re.search(r'^(?:Q\d+\.?|Question(?:\s*\d+)?:?|Objection(?:\s*\d+)?:?|Scenario(?:\s*\d+)?:?|Client(?:\s*says)?:?)\s*(.+?)\n+(?:Answer:?|Pitch:?|Response:?|Strategy:?|Solution:?)\s*\n+(.+)', block, re.DOTALL | re.IGNORECASE)
                    if sub_m2:
                        q_text = sub_m2.group(1).strip()
                        pitch_text = sub_m2.group(2).strip()
                        context_text = f"Custom strategy for objection: {q_text[:80]}"
                    else:
                        lines = [line.strip() for line in block.split("\n") if line.strip()]
                        if len(lines) >= 2:
                            q_text = re.sub(r'^(?:Q\d+\.?|Question(?:\s*\d+)?:?|Objection(?:\s*\d+)?:?|Scenario(?:\s*\d+)?:?|Client(?:\s*says)?:?)\s*', '', lines[0], flags=re.IGNORECASE).strip()
                            pitch_text = "\n".join(lines[1:]).strip()
                            context_text = f"Sales playbook strategy from {source_filename}"
                        else:
                            continue

                if q_text and pitch_text:
                    cards.append({
                        "q_number": card_id,
                        "question": sanitize_unicode(q_text),
                        "context": sanitize_unicode(context_text),
                        "pitch": sanitize_unicode(pitch_text),
                        "source": sanitize_unicode(source_filename)
                    })
                    card_id += 1

        # 2. Numbered Items / Rules / Chapters / Strategies (e.g. 1. , 2. , 1) , Rule 1: , Chapter 1:)
        if len(cards) < 2:
            num_pattern = r'(?=(?:^|\n)\s*(?:\d+[\.\)]\s+|Rule\s*\d+:?|Strategy\s*\d+:?|Tip\s*\d+:?|Step\s*\d+:?|Section\s*\d+:?|Chapter\s*\d+:?|Module\s*\d+:?|Topic\s*\d+:?))'
            num_blocks = [b.strip() for b in re.split(num_pattern, clean_text, flags=re.IGNORECASE) if b.strip() and len(b.strip()) >= 10]
            if len(num_blocks) >= 2:
                card_id = 1
                for nb in num_blocks:
                    lines = [l.strip() for l in nb.split("\n") if l.strip()]
                    if lines:
                        title = re.sub(r'^(?:\d+[\.\)]\s+|Rule\s*\d+:?|Strategy\s*\d+:?|Tip\s*\d+:?|Step\s*\d+:?|Section\s*\d+:?|Chapter\s*\d+:?|Module\s*\d+:?|Topic\s*\d+:?)\s*', '', lines[0], flags=re.IGNORECASE).strip()
                        body = "\n".join(lines[1:]).strip() if len(lines) > 1 else nb
                        if not title:
                            title = f"Strategy Rule #{card_id}"
                        cards.append({
                            "q_number": card_id,
                            "question": sanitize_unicode(title),
                            "context": f"Playbook rule #{card_id} from {source_filename}",
                            "pitch": sanitize_unicode(body),
                            "source": sanitize_unicode(source_filename)
                        })
                        card_id += 1

        # 3. Markdown Headers (# Header, ## Section)
        if len(cards) < 2:
            header_blocks = re.split(r'(?:^|\n)(?=#{1,4}\s+)', clean_text)
            header_blocks = [hb.strip() for hb in header_blocks if hb.strip() and len(hb.strip()) >= 10]
            if len(header_blocks) >= 2:
                card_id = 1
                for hb in header_blocks:
                    lines = [l.strip() for l in hb.split("\n") if l.strip()]
                    if lines:
                        header = re.sub(r'^#{1,4}\s*', '', lines[0]).strip()
                        body = "\n".join(lines[1:]).strip() if len(lines) > 1 else hb
                        cards.append({
                            "q_number": card_id,
                            "question": sanitize_unicode(header),
                            "context": f"Section: {header} from {source_filename}",
                            "pitch": sanitize_unicode(body),
                            "source": sanitize_unicode(source_filename)
                        })
                        card_id += 1

        # 4. Bullet Points & List Items (• , - , * , ▪ , ►)
        if len(cards) < 2:
            bullet_pattern = r'(?:^|\n)\s*[•\-\*▪►–]\s+(.+)'
            bullet_matches = re.findall(bullet_pattern, clean_text)
            if len(bullet_matches) >= 3:
                card_id = 1
                for bm in bullet_matches:
                    bm_clean = bm.strip()
                    if len(bm_clean) > 25:
                        parts = bm_clean.split(":", 1)
                        if len(parts) == 2 and len(parts[0]) < 60:
                            q_text = parts[0].strip()
                            p_text = parts[1].strip()
                        else:
                            first_sentence_m = re.search(r'^([^.!?]+[.!?])', bm_clean)
                            q_text = first_sentence_m.group(1).strip() if first_sentence_m else bm_clean[:50] + "..."
                            p_text = bm_clean
                        cards.append({
                            "q_number": card_id,
                            "question": sanitize_unicode(q_text),
                            "context": f"Bullet playbook item #{card_id} from {source_filename}",
                            "pitch": sanitize_unicode(p_text),
                            "source": sanitize_unicode(source_filename)
                        })
                        card_id += 1

        # 5. Universal High-Resolution Sliding Semantic Window Chunker (80-140 words per chunk)
        if len(cards) < 2:
            words = clean_text.split()
            if len(words) > 40:
                target_chunk_size = 90
                overlap = 20
                i = 0
                card_id = 1
                while i < len(words):
                    chunk_words = words[i:i + target_chunk_size]
                    chunk_text = " ".join(chunk_words).strip()
                    first_sent_m = re.search(r'^([^.!?\n]+[.!?])', chunk_text)
                    if first_sent_m and len(first_sent_m.group(1).strip()) > 15:
                        title = first_sent_m.group(1).strip()
                    else:
                        title = " ".join(chunk_words[:7]).strip() + "..."

                    cards.append({
                        "q_number": card_id,
                        "question": sanitize_unicode(f"Strategy #{card_id}: {title}"),
                        "context": f"Playbook semantic excerpt #{card_id} from {source_filename}",
                        "pitch": sanitize_unicode(chunk_text),
                        "source": sanitize_unicode(source_filename)
                    })
                    card_id += 1
                    if i + target_chunk_size >= len(words):
                        break
                    i += (target_chunk_size - overlap)

        # Fallback if text is very short (<40 words)
        if not cards:
            cards.append({
                "q_number": 1,
                "question": sanitize_unicode(f"Playbook Summary: {source_filename}"),
                "context": f"Uploaded document {source_filename}",
                "pitch": sanitize_unicode(clean_text[:1200]),
                "source": sanitize_unicode(source_filename)
            })

        logger.info(f"Chunked '{source_filename}' into {len(cards)} strategy battlecards.")
        return cards
